import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

import pytest
import tempfile
from core.parsers import get_parser
from core.parsers.base_parser import _REGISTRY


def test_txt_parser_utf8(tmp_path):
    f = tmp_path / "test.txt"
    f.write_text("心脏解剖\n\n左心室壁最厚", encoding="utf-8")
    parser = get_parser(f)
    result = parser.parse(f)
    assert "心脏解剖" in result
    assert "左心室壁" in result


def test_txt_parser_gbk(tmp_path):
    f = tmp_path / "test.txt"
    f.write_bytes("心脏解剖".encode("gbk"))
    parser = get_parser(f)
    result = parser.parse(f)
    assert "心脏" in result


def test_unsupported_format_raises():
    p = Path("test.xyz")
    with pytest.raises(ValueError, match="不支持的文件格式"):
        get_parser(p)


def test_all_parsers_registered():
    assert ".pdf" in _REGISTRY
    assert ".docx" in _REGISTRY
    assert ".txt" in _REGISTRY


def test_docx_parser(tmp_path):
    from docx import Document as DocxDocument
    f = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("心脏有四个腔")
    doc.add_paragraph("左心房、左心室、右心房、右心室")
    doc.save(str(f))
    parser = get_parser(f)
    result = parser.parse(f)
    assert "心脏有四个腔" in result
